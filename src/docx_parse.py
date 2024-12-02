from docx import Document

class WordDocuments:

    def __init__(self, file_path):
        self.document = Document(file_path)

    def get_all_text(self):
        all_text = ""

        paragraphs = self.__get_paragraphs()
        tables = self.__get_tables()

        all_text = paragraphs + tables

        return all_text

    def __get_paragraphs(self):
        para_text = ""

        for paragraph in self.document.paragraphs:
            para_text += paragraph.text + '\n'
  
        return para_text

    def __get_tables(self):
        table_text = ""

        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        table_text += para.text + " "

        return table_text